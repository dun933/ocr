import os

os.environ["CUDA_VISIBLE_DEVICES"] = "2"
import tensorflow as tf
from keras.backend.tensorflow_backend import set_session

config = tf.ConfigProto()
config.gpu_options.per_process_gpu_memory_fraction = 0.01
set_session(tf.Session(config=config))

import torch
from pdf2word.test import pdf_is_text, extract_pdf_text

# from Rec_text import rec_txt
from pdf2image import convert_from_path

# from ctpn.ctpn_blstm_test_full import text_predict
# from densent_ocr.model import predict
from pan.predict import text_predict
from crnn_torch.model import predict
from extract_table_scale import extract_table, generate_table
import cv2
#
import numpy as np
import docx
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from PIL import Image
import fitz
from docx.oxml.ns import qn
import time
from sanic import Sanic, response
from layout import sort_paragraph, save2docx
from direction_detection.SkewDetect import skew_detect

from crnn_seq2seq_ocr.inference import attention
from viterbi import calculate


a = Image.open('test.png').convert('RGB')
ori_w, ori_h = a.size
b = a.resize((1000, 1000))
scale_w, scale_h = b.size
b = np.array(b)
scale_w, scale_h = ori_w / scale_w, ori_h / scale_h
text_predict(b, scale_w, scale_h, a)

app = Sanic(__name__)


def single_ocr(document, img_name, start_page, new_url):
    img_name = skew_detect.get_rotated_img(img_name)

    ori_img = np.array(img_name)
    ori_w, ori_h = img_name.size
    img_name.thumbnail((1500, 1500), Image.ANTIALIAS)
    # img_name = img_name.resize((1500, 1500), Image.ANTIALIAS)
    # img_name = img_name.convert('RGB')
    scale_w, scale_h = img_name.size
    # print(scale_w, scale_h)
    scale_w, scale_h = ori_w / scale_w, ori_h / scale_h
    print('原图大小:', ori_w, ori_h, '缩放比例:', scale_w, scale_h)
    img = np.array(img_name)
    # B_channel, G_channel, R_channel = cv2.split(img)
    # cv2.imwrite('test.png', R_channel)
    # img = cv2.cvtColor(R_channel, cv2.COLOR_GRAY2BGR)
    start = time.time()
    images = text_predict(img, scale_w, scale_h, ori_img)
    torch.cuda.empty_cache()
    print('ctpn time: ', time.time() - start)
    # new_images = []
    # images = new_images
    # Image.fromarray(img).save('paragraph.jpg')
    # Image.fromarray(img).show()
    try:
        tables = extract_table(ori_img)
        if tables == 'not table':
            has_table = False
        else:
            has_table = True
            # for table in tables:
            #     table[0][1][1] = table[0][1][1] / scale_h
            #     table[0][1][3] = table[0][1][3] / scale_h
    except:
        has_table = False
    print(2222222222222222222222222, has_table)
    results = []
    start = time.time()
    for index, j in enumerate(images):
        # if j[1].any() and j[1].shape[0] < j[1].shape[1] * 1.5:
        try:
            if has_table:
                count = 0
                for table in tables:
                    if table[0][1][1] + table[0][1][3] > j[0][1] > table[0][1][1]:
                        continue
                    else:
                        count += 1
                if count == len(tables):
                    content = predict(Image.fromarray(j[1]).convert('L'))
                    ori_content = [i[0] for i in content[0]]
                    prob_content = [[i, j] for i, j in zip(content[0], content[1])]
                    for indexi, i in enumerate(content[1]):
                        if i[0] > 0.9:
                            content[0][indexi] = content[0][indexi][0]
                            content[1][indexi] = [-1]
                    while 1:
                        try:
                            content[1].remove([-1])
                        except:
                            break
                    # ori_content = [i[0] for i in content[0]]
                    # with open(os.path.splitext(new_url)[0] + '.txt', 'a', encoding='utf-8') as f:
                    # for index, i in enumerate(content[1]):
                    #     if i[0] > 0.9:
                    #         content[0][index] = content[0][index][0]
                    #         content[1].pop(index)
                    # if i[0] < 0.9:
                    #     img = Image.fromarray(j[1]).convert('L')
                    #     width, height = img.size[0], img.size[1]
                    #     scale = height * 1.0 / 32
                    #     width = int(width / scale)
                    #
                    #     img = img.resize([width, 32], Image.ANTIALIAS)
                    #     img = np.array(img)
                    #     new_img = img[:, (content[2][index] - 1) * 8:(content[2][index] + 2) * 8]
                    #     word, prob = attention(new_img)
                    #     if prob > 0.9:
                    #         content[0][index] = word[0]
                    #         content[1].pop(index)
                    # else:
                    #     content[0][index] = content[0][index][0]
                    #     content[1].pop(index)
                    content = calculate(content)
                    # for i, j_i in zip(ori_content, content):
                    #     if j_i != i:
                    #         f.write(i + '------------>' + j_i + '\n')
                    # content = rec_txt(j[1])
                    # torch.cuda.empty_cache()
                    results.append([j[0], content.replace('“', '').replace('‘', '')])
            else:
                content = predict(Image.fromarray(j[1]).convert('L'))
                ori_content = [i[0] for i in content[0]]
                prob_content = [[i, j] for i, j in zip(content[0], content[1])]
                for indexi, i in enumerate(content[1]):
                    if i[0] > 0.9:
                        content[0][indexi] = content[0][indexi][0]
                        content[1][indexi] = [-1]
                while 1:
                    try:
                        content[1].remove([-1])
                    except:
                        break
                # ori_content = [i[0] for i in content[0]]
                # with open(os.path.splitext(new_url)[0] + '.txt', 'a', encoding='utf-8') as f:
                # for index, i in enumerate(content[1]):
                #     if i[0] > 0.9:
                #         content[0][index] = content[0][index][0]
                #         content[1].pop(index)
                # if i[0] < 0.9:
                #     img = Image.fromarray(j[1]).convert('L')
                #     width, height = img.size[0], img.size[1]
                #     scale = height * 1.0 / 32
                #     width = int(width / scale)
                #
                #     img = img.resize([width, 32], Image.ANTIALIAS)
                #     img = np.array(img)
                #     new_img = img[:, (content[2][index] - 1) * 8:(content[2][index] + 2) * 8]
                #     word, prob = attention(new_img)
                #     if prob > 0.9:
                #         content[0][index] = word[0]
                #         content[1].pop(index)
                # else:
                #     content[0][index] = content[0][index][0]
                #     content[1].pop(index)
                content = calculate(content)
                # for i, j_i in zip(ori_content, content):
                #     if j_i != i:
                #         f.write(i + '------------>' + j_i + '\n')
                # content = rec_txt(j[1])
                # torch.cuda.empty_cache()
                results.append([j[0], content.replace('“', '').replace('‘', '')])
        except Exception as e:
            print(e)
            continue
    torch.cuda.empty_cache()
    print(33333333333333333, time.time() - start)
    results = sorted(results, key=lambda i: i[0][1])
    new_results = results
    line_images = []
    cut_index = 0
    curr_index = 0
    for index, i in enumerate(new_results):
        if index == len(new_results) - 1:
            if cut_index < index:
                line_images.append(new_results[cut_index:index])
                line_images.append(new_results[index:])
            else:
                line_images.append(new_results[index:])
            break
        # if abs(new_results[index + 1][0][1] - new_results[index][0][1]) > (
        #         new_results[index][0][7] - new_results[index][0][1]) * 4 / 5:
        #     line_images.append(new_results[cut_index: index + 1])
        #     cut_index = index + 1
        if abs(new_results[index + 1][0][1] - new_results[curr_index][0][1]) > (
                new_results[curr_index][0][7] - new_results[curr_index][0][1]) * 4 / 5:
            line_images.append(new_results[cut_index: index + 1])
            cut_index = index + 1
            curr_index = index + 1
    for index, i in enumerate(line_images):
        line_images[index] = sorted(i, key=lambda a: a[0][0])
    texts = []
    position = []
    for i in line_images:
        text = ''
        for index, j in enumerate(i):
            try:
                position.append([j[0], j[1]])
                if index == len(i) - 1:
                    text += j[1]
                elif abs(i[index + 1][0][0] - i[index][0][6]) > 3 * (
                        abs(i[index][0][6] - i[index][0][0]) / len(i[index][1])):
                    text += j[1] + ' '
                else:
                    text += j[1]
            except:
                continue
        texts.append([[i[0][0], i[-1][0]], text])
    print(img_name.size)
    if has_table:
        for table in tables:
            table_index = 0
            for index, i in enumerate(texts):
                # print(i)
                # print(type(i[0]), type(table[1][1]))
                if i[0] == 'table':
                    # print(table[0][1])
                    if table[0][1][1] > i[1][3][1]:
                        table_index = index + 1
                elif table[0][1][1] > i[0][0][1]:
                    table_index = index + 1
            try:
                texts.insert(table_index, ['table', generate_table(table, ori_img)])
            except Exception as e:
                print(e)
                continue
    # import pickle
    # pickle.dump(texts, open('texts.pkl', 'wb'))
    try:
        texts = sort_paragraph(Image.fromarray(ori_img), texts)
    except Exception as e:
        print(e)
        return document, position
    document = save2docx(document, texts, Image.fromarray(ori_img), start_page)
    return document, position


@app.route('/full', methods=['POST'])
def full_ocr(request):
    print('start full ocr')
    start = time.time()
    pdf_file = list(request.json.get('input')[0].keys())[0]
    new_url = list(request.json.get('input')[0].values())[0]
    start_page = request.json.get('start_page')
    end_page = request.json.get('end_page')
    flag = request.json.get('flag')
    positions = []
    pdf_images = []
    print(pdf_file)
    page_num = 1
    try:
        if pdf_file.lower().endswith('.pdf'):
            if pdf_is_text(pdf_file):
                document = docx.Document()
                document.styles['Normal'].font.name = u'宋体'
                document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                for i in range(int(start_page) - 1, int(end_page)):
                    p = document.add_paragraph()
                    run = p.add_run('[W-{}-L]'.format(i))
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    result, img = extract_pdf_text(pdf_file, i)
                    document = save2docx(document, result, img, int(start_page))
                    if i == int(end_page) - 1:
                        continue
                    document.add_section()
                if flag == 'ddwork':
                    files = os.path.splitext(pdf_file)[0] + '_{}_{}'.format(start_page, end_page) + '.docx'
                else:
                    files = os.path.splitext(pdf_file)[0] + '.docx'
                new_url = new_url + files.split('/')[-1]
                document.save(new_url)
                return response.json(
                    {'result': 'true', 'Documents': [{'new_url': new_url}], 'message': '请求成功', 'positions': []})

            # if start_page == end_page:
            #     pdf = fitz.open(pdf_file)
            #     page = pdf[int(start_page) - 1]
            #     trans = fitz.Matrix(3, 3).preRotate(0)
            #     pm = page.getPixmap(matrix=trans, alpha=False)
            #     img = Image.frombytes("RGB", [pm.width, pm.height], pm.samples)
            #     pdf_images.append(img)
            #     page_num = start_page
            # else:
            pdf = fitz.open(pdf_file)
            for i in range(int(start_page) - 1, int(end_page)):
                page = pdf[i]
                trans = fitz.Matrix(3, 3).preRotate(0)
                pm = page.getPixmap(matrix=trans, alpha=False)
                img = Image.frombytes("RGB", [pm.width, pm.height], pm.samples)
                pdf_images.append(img)
            page_num = start_page
        else:
            page_num = start_page
            pdf_images = [Image.open(pdf_file).convert('RGB')]
        document = docx.Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        if flag == 'ddwork':
            files = os.path.splitext(pdf_file)[0] + '_{}_{}'.format(start_page, end_page) + '.docx'
        else:
            files = os.path.splitext(pdf_file)[0] + '.docx'
        new_url = new_url + files.split('/')[-1]
        for index, img_name in enumerate(pdf_images):
            print(index)
            if time.time() - start > 40 * len(pdf_images):
                return response.json({'result': 'false', 'Documents': [], 'message': '请求超时'})
            else:
                p = document.add_paragraph()
                run = p.add_run('[W-{}-L]'.format(index + int(page_num)))
                run.font.color.rgb = RGBColor(255, 255, 255)
                document, position = single_ocr(document, img_name, int(start_page) + index, new_url)
                positions = position
            if index == len(pdf_images) - 1:
                continue
            document.add_section()
        document.save(new_url)
    except Exception as e:
        print('error', e)
        return response.json({'result': 'false', 'Documents': [], 'message': '请求失败'})
    return response.json(
        {'result': 'true', 'Documents': [{'new_url': new_url}], 'message': '请求成功', 'positions': positions})


@app.route('/check')
def health_check(request):
    return response.text('ok')


if __name__ == '__main__':
    app.config.KEEP_ALIVE = False
    app.config.RESPONSE_TIMEOUT = 7200
    app.run(host='0.0.0.0', port=8123)
