import os
import torch
from pdf2word.test import pdf_is_text, extract_pdf_text

import cv2
import requests
import base64
#
import numpy as np
import docx
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from PIL import Image
import fitz
from docx.oxml.ns import qn
import time
from sanic import Sanic, response
from direction_detection.SkewDetect import skew_detect
from layout_unet.unet_parse import unetParse
from viterbi import calculate
from config_url import DETECT_URL, RECOGNISE_URL


# DETECT_URL = 'http://172.30.81.191:32021/text_predict'
# RECOGNISE_URL = 'http://172.30.81.191:32020/predict'


def single_ocr(document, img_name, start_page, new_url):
    start2 = time.time()
    img_name = skew_detect.get_rotated_img(img_name)

    ori_img = np.array(img_name)
    _, ori = cv2.imencode('.jpg', ori_img)
    ori = base64.b64encode(ori.tostring())

    ori_w, ori_h = img_name.size
    print('旋转时间：', time.time() - start2)
    while_i = 0
    start = time.time()
    images = []
    while 1:
        img_name.thumbnail((1500 - while_i * 100, 1500 - while_i * 100), Image.ANTIALIAS)
        # img_name = img_name.resize((1500, 1500), Image.ANTIALIAS)
        # img_name = img_name.convert('RGB')
        scale_w, scale_h = img_name.size
        # print(scale_w, scale_h)
        scale_w, scale_h = ori_w / scale_w, ori_h / scale_h
        print('原图大小:', ori_w, ori_h, '缩放比例:', scale_w, scale_h)
        img = np.array(img_name)
        # B_channel, G_channel, R_channel = cv2.split(img)
        # cv2.imwrite('test.png', R_channel)
        # img = cv2.cvtColor(R_channel, cv2.COLOR_GRAY2BGR)
        _, img = cv2.imencode('.jpg', img)
        img = base64.b64encode(img.tostring())
        data = {'img': img, 'scale_w': scale_w, 'scale_h': scale_h, 'ori_img': ori}
        images_json = requests.post(DETECT_URL, data=data)
        # images = text_predict(img, scale_w, scale_h, ori_img)
        torch.cuda.empty_cache()
        while_i += 1
        if images_json.json() != '':
            for i in images_json.json():
                image = base64.b64decode(i[1])
                image = np.fromstring(image, np.uint8)
                image = cv2.imdecode(image, cv2.IMREAD_COLOR)
                images.append([i[0], image])
            break
    print('ctpn time: ', time.time() - start)

    results = []
    start = time.time()
    # print(1111111, images)
    for index, j in enumerate(images):
        # if j[1].any() and j[1].shape[0] < j[1].shape[1] * 1.5:
        try:

            _, img = cv2.imencode('.jpg', j[1])
            img = base64.b64encode(img.tostring())
            data = {'img': img}
            content = requests.post(RECOGNISE_URL, data=data).json()[:2]
            # ori_content = [i[0] for i in content[0]]
            # prob_content = [[i, j] for i, j in zip(content[0], content[1])]
            for indexi, i in enumerate(content[1]):
                if i[0] > 0.9:
                    content[0][indexi] = content[0][indexi][0]
                    content[1][indexi] = [-1]
            while 1:
                try:
                    content[1].remove([-1])
                except:
                    break
            content = calculate(content)

            results.append([j[0], content.replace('“', '').replace('‘', '')])
        except Exception as e:
            print('h w', e)
            continue
    print('识别时间', time.time() - start)
    print(results)
    start = time.time()
    torch.cuda.empty_cache()
    results = sorted(results, key=lambda i: i[0][1])
    new_results = results
    line_images = []
    cut_index = 0
    curr_index = 0
    for index, i in enumerate(new_results):
        if index == len(new_results) - 1:
            if cut_index < index:
                line_images.append(new_results[cut_index:index])
                line_images.append(new_results[index:])
            else:
                line_images.append(new_results[index:])
            break
        # if abs(new_results[index + 1][0][1] - new_results[index][0][1]) > (
        #         new_results[index][0][7] - new_results[index][0][1]) * 4 / 5:
        #     line_images.append(new_results[cut_index: index + 1])
        #     cut_index = index + 1
        if abs(new_results[index + 1][0][1] - new_results[curr_index][0][1]) > (
                new_results[curr_index][0][7] - new_results[curr_index][0][1]) * 4 / 5:
            line_images.append(new_results[cut_index: index + 1])
            cut_index = index + 1
            curr_index = index + 1
    for index, i in enumerate(line_images):
        line_images[index] = sorted(i, key=lambda a: a[0][0])
    texts = []
    position = []
    for i in line_images:
        text = ''
        for index, j in enumerate(i):
            try:
                position.append([j[0], j[1]])
                if index == len(i) - 1:
                    text += j[1]
                elif abs(i[index + 1][0][0] - i[index][0][6]) > 3 * (
                        abs(i[index][0][6] - i[index][0][0]) / len(i[index][1])):
                    text += j[1] + ' '
                else:
                    text += j[1]
            except:
                continue
        texts.append([[i[0][0], i[-1][0]], text])
    print(img_name.size)

    document = unetParse(ori_img, texts, document).save2docx()
    # try:
    #     texts = sort_paragraph(Image.fromarray(ori_img), texts)
    # except Exception as e:
    #     print(e)
    #     return document, position
    # document = save2docx(document, texts, Image.fromarray(ori_img), start_page)
    print('版式表格时间：', time.time() - start)
    return document, position


def full_ocr(request):
    print('start full ocr')
    start = time.time()
    start2 = time.time()
    pdf_file = list(request.json.get('input')[0].keys())[0]
    new_url = list(request.json.get('input')[0].values())[0]
    start_page = request.json.get('start_page')
    end_page = request.json.get('end_page')
    print('page', start_page, end_page)
    flag = request.json.get('flag')

    positions = []
    pdf_images = []
    print(pdf_file)
    page_num = 1
    try:
        if pdf_file.lower().endswith('.pdf'):
            if pdf_is_text(pdf_file):
                document = docx.Document()
                document.styles['Normal'].font.name = u'宋体'
                document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                for i in range(int(start_page) - 1, int(end_page)):
                    p = document.add_paragraph()
                    run = p.add_run('[W-{}-L]'.format(i))
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    result, img = extract_pdf_text(pdf_file, i)
                    document = save2docx(document, result, img, int(start_page))
                    if i == int(end_page) - 1:
                        continue
                    document.add_section()
                if flag == 'ddwork':
                    files = os.path.splitext(pdf_file)[0] + '_{}_{}'.format(start_page, end_page) + '.docx'
                else:
                    files = os.path.splitext(pdf_file)[0] + '.docx'
                new_url = new_url + files.split('/')[-1]
                document.save(new_url)
                return response.json(
                    {'result': 'true', 'Documents': [{'new_url': new_url}], 'message': '请求成功', 'positions': []})

            # if start_page == end_page:
            #     pdf = fitz.open(pdf_file)
            #     page = pdf[int(start_page) - 1]
            #     trans = fitz.Matrix(3, 3).preRotate(0)
            #     pm = page.getPixmap(matrix=trans, alpha=False)
            #     img = Image.frombytes("RGB", [pm.width, pm.height], pm.samples)
            #     pdf_images.append(img)
            #     page_num = start_page
            # else:
            pdf = fitz.open(pdf_file)
            for i in range(int(start_page) - 1, int(end_page)):
                page = pdf[i]
                trans = fitz.Matrix(3, 3).preRotate(0)
                pm = page.getPixmap(matrix=trans, alpha=False)
                img = Image.frombytes("RGB", [pm.width, pm.height], pm.samples)
                pdf_images.append(img)
            pdf.close()
            page_num = start_page
        else:
            page_num = start_page
            pdf_images = [Image.open(pdf_file).convert('RGB')]
        document = docx.Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        print('读取时间：', time.time() - start2)
        if flag == 'ddwork':
            files = os.path.splitext(pdf_file)[0] + '_{}_{}'.format(start_page, end_page) + '.docx'
        else:
            files = os.path.splitext(pdf_file)[0] + '.docx'
        new_url = new_url + files.split('/')[-1]
        for index, img_name in enumerate(pdf_images):
            print(index)
            if time.time() - start > 40 * len(pdf_images):
                return response.json({'result': 'false', 'Documents': [], 'message': '请求超时'})
            else:
                start1 = time.time()
                p = document.add_paragraph()
                run = p.add_run('[W-{}-L]'.format(index + int(page_num)))
                run.font.color.rgb = RGBColor(255, 255, 255)
                print('docx处理时间：', time.time() - start1)
                document, position = single_ocr(document, img_name, int(start_page) + index, new_url)
                positions = position
            if index == len(pdf_images) - 1:
                continue
            document.add_section()
        document.save(new_url)
        print('总时间:', time.time() - start)
    except Exception as e:
        print('error', e)
        return response.json({'result': 'false', 'Documents': [], 'message': '请求失败'})
    return response.json(
        {'result': 'true', 'Documents': [{'new_url': new_url}], 'message': '请求成功', 'positions': positions})
