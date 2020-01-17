import math
import os
import re

# from font import get_detail

from PIL import Image
import numpy as np
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from add_picture_float import add_picture_float
from restore_table import restore_table


def save2docx(doc, result, img, jpg_page):

    sec = doc.sections[-1]
    header = sec.header
    header.is_linked_to_previous = False
    footer = sec.footer
    footer.is_linked_to_previous = False





    # for index, i in enumerate(result):
    #     if i[1].startswith('header'):
    #         header.paragraphs[-1].style.font.size = Pt(9)
    #         if i[1].endswith('separate'):
    #             header.paragraphs[-1].text = i[0][0] + '\t' + i[0][1]
    #             header.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    #             continue
    #         if i[1].endswith('center'):
    #             header.paragraphs[-1].text = i[0]
    #             header.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #             continue
    #         header.paragraphs[-1].text = i[0]
    #         header.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    #     elif i[1].startswith('footer'):
    #         footer.paragraphs[-1].style.font.size = Pt(8)
    #         footer.paragraphs[-1].text = i[0]
    #         footer.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #     elif i[1] == 'table':
    #         doc = restore_table(doc, i, img)
    #         if index == len(result)-1 or index == len(result)-2:
    #             doc.add_paragraph()
    #     else:
    #         if i[1] == 'center':
    #             paragraph = doc.add_paragraph()
    #             paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #             paragraph.style.font.size = Pt(i[2])  #
    #             runner = paragraph.add_run(i[0])
    #             runner.font.size = Pt(i[2])
    #             runner.bold = True
    #         elif i[1] == 'right':
    #             paragraph = doc.add_paragraph()
    #             paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    #             paragraph.style.font.size = Pt(i[2])
    #             runner = paragraph.add_run(i[0])
    #             runner.font.size = Pt(i[2])
    #         else:
    #             paragraph = doc.add_paragraph()
    #             paragraph.paragraph_format.first_line_indent = Inches(0.3)
    #             paragraph.style.font.size = Pt(i[2])
    #             runner = paragraph.add_run(i[0])
    #             runner.font.size = Pt(i[2])
    #         if i[3] > 0:
    #             for j in range(i[3]):
    #                 blank_paragraph = doc.add_paragraph('')
    #                 blank_paragraph.style.font.size = Pt(7)
    #                 blank_paragraph.paragraph_format.space_after = Pt(0)
    #                 blank_paragraph.paragraph_format.line_spacing_rule = Pt(0)
    #         if i[4] == 0:
    #             paragraph.paragraph_format.line_spacing_rule = 0
    #         if i[3] > 0:
    #             space_after_code = 0
    #         paragraph.paragraph_format.space_after = Pt(space_after_code)
    #         paragraph.paragraph_format.line_spacing_rule = line_spacing_rule_code  # 1
    #
    # # 写入图片
    # if img_pos_list:
    #     data = [i for i in result if i[1] in ['general', 'sub', 'table', 'center']]
    #     if not data:
    #         print('result为空')
    #         doc.add_paragraph()
    #     for i in img_pos_list:
    #         try:
    #             Image.fromarray(np.array(img)[i[0][1]:i[0][3], i[0][0]:i[0][2]]).save('img_s.jpg')
    #             if i[0][0] > img.width / 2:
    #                 pos_w = 7534275 * (i[0][0] / img.width)
    #             else:
    #                 pos_w = 7534275 * (i[0][0] / img.width)
    #             if i[0][1] < img.width / 2:
    #                 pos_y = 10592435 * (i[0][1] / img.height)
    #             else:
    #                 pos_y = 10592435 * (2 * i[0][1] - i[0][3]) / img.height
    #             pos = f'{i[0][0]}_{i[0][1]}_{i[0][2]}_{i[0][3]}_{jpg_page}_{img.width}_{img.height}'
    #             add_picture_float(doc, 'img_s.jpg', (pos_w, pos_y), pos, scales=5)
    #             os.remove('img_s.jpg')
    #         except:
    #             pass

    # doc.add_section()
    return doc