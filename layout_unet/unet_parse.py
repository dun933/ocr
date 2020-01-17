import os
import cv2
from PIL import Image
import numpy as np
import json
import pickle
import re

import requests
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.oxml.ns import qn
import base64
import io


from restore_table import restore_table
from layout_unet.to_word import generate_table, extract_table
from layout_unet.post_process import thresholding, find_polygonal_regions
from config_url import LAYOUT_URL, TABLE_URL

# LAYOUT_URL = 'http://172.30.20.154:32081/layout'
# TABLE_URL = 'http://172.30.20.154:32082/table_line'
COLOR_DICT = np.array([[255, 255, 255], [1, 1, 1], [64, 64, 64], [100, 100, 100], [150, 150, 150], [200, 200, 200]])


# 判断行是否在表格区域
def checkIsInTable(tables, texts):
    # print(tables)
    # print(texts)
    for table in tables:
        table_bbox = table
        if table[0] < texts[0][0] + 0.5*texts[0][1] < table[0]+table[2] and table[1] < texts[0][1] + 0.5*texts[0][3] < table[1]+table[3]:
            return True
    return False


class unetParse(object):
    def __init__(self, image, line_infos, document):
        self.image = image
        self.img_dict = self.unetpic2typepos()
        self.line_infos = line_infos
        self.line_lack = None
        self.result_temp = None
        self.lines = []  # 元素格式示例 ['第1页共2页', 10.5, 2353, 'footer_center'] [文字， 字体大小， 纵坐标， 居中格式]
        self.document = document
        self.tables = []

    def unetpic2typepos(self):
        retval, buffer = cv2.imencode('.jpg', self.image)
        pic_str = base64.b64encode(buffer)
        pic_str = pic_str.decode()
        res = requests.post(LAYOUT_URL, data={'image': pic_str})
        print(res.content)
        img = self.image.copy()
        for k, v in res.json().items():
            for i in v:
                cv2.putText(img, k, (i[0][0], i[0][1]), cv2.FONT_HERSHEY_COMPLEX, 1, (0,0,255), 10)
                cv2.rectangle(img, (i[0][0], i[0][1]), (i[0][0]+i[0][2], i[0][1]+i[0][3]), (255, 0, 0), 1)
        Image.fromarray(img).save('layout_unet/layout.jpg')
        return res.json()


    def get_font(self, line_pos):
        # 获取当前行的字体大小
        h = line_pos[5]-line_pos[1]
        if h < 30:
            return 9
        elif 30 <= h <= 50:
            return 10.5
        elif 50 < h <= 65:
            return 12
        elif 65 < h < 80:
            return 14
        return 16


    def layout(self, texts, texts_w, line_type):
        # texts 为在此区域的检测行
        # im_w 为区域的宽度

        # 将框选区域移动到最左端
        # print(texts)
        # if line_type == 'text':
        #     line_x_list = [texts[0][0]]
        #     for text in texts[1]:
        #         for line_bbox in text[0]:
        #             # print(2,line_bbox)
        #             for i in range(len(line_bbox)):
        #                 if i in [0, 2, 4, 6]:
        #                     line_x_list.append(line_bbox[i])
            
        #     min_x = min(line_x_list)
        #     for index in range(len(texts[1])):
        #         line = texts[1][index]
        #         temp = [0, 2, 4, 6]
        #         texts[1][index] = [[[line[0][0][x]-min_x if x in temp else line[0][0][x] for x in range(len(line[0][0]))],
        #         [line[0][0][x]-min_x if x in temp else line[0][0][x] for x in range(len(line[0][0]))]], line[1]]
            
        new_texts = []
        w = self.image.shape[1]
        for i in texts:
            if len(texts) == 1:
                # 考虑只有一行时的文本居中、居左、居右问题
                # print('texts_len', i)
                if i[0][0][0] > w*0.5:
                    new_texts.append([i[0], i[1], 'right'])
                elif i[0][0][2] < w*0.5:
                    new_texts.append([i[0], i[1], 'left'])
                elif i[0][-1][6] < 0.80 * w and i[0][0][0] > 0.20 * w and 0.9 * w < i[0][0][0] + i[0][-1][6] < 1.1 * w:
                    new_texts.append([i[0], i[1], 'center'])
                else:
                    new_texts.append([i[0], i[1], 'sub'])
                break
            end_flg = 'general'
            if re.compile('(^([0-9]{1,2}|[一二三四五六七八九十]{1,3})[、.])|(^[A-Z][）)])').search(i[1]) or (
                re.search(r'(^第.{1,3}条)|(^[(（][0-9一二三四五六七八九十]{1,3}[）)])', i[1])):
                if len(new_texts) > 0:
                    new_texts[-1][-1] = 'sub'
            if i[0][0][2] < w * 0.7:
                end_flg = 'sub'
            new_texts.append([i[0], i[1], end_flg])
        # print(111, new_texts)
        document = ''
        text_tmp = ['', 7.5]
        paragraph_type = 'text'
        for i in new_texts:
            end_flg = ''
            text_tmp[0] += i[1]
            # print(len(new_texts), i)
            # if len(new_texts) == 1:
            #     print('new', i)
            #     text_tmp[1] = self.get_font(i[0][0])
            #     if i[-1] != 'sub':
            #         paragraph_type = f'{line_type}_{i[-1]}'
            #         print(123, paragraph_type)
                
            #     text_tmp.append(paragraph_type)
            #     # print(222, text_tmp)
            #     self.lines.append(text_tmp)
            #     break
            # if i[-1]
            if i[-1] in ['right', 'left', 'center']:
                paragraph_type = f'{line_type}_{i[-1]}'
                text_tmp[1] = self.get_font(i[0][0])
                text_tmp.extend(([i[0][0][1], paragraph_type]))
                self.lines.append(text_tmp)
                text_tmp = ['', 0]
                continue
            if i[-1] == 'sub':
                if line_type != 'text':
                    paragraph_type = line_type
                text_tmp[1] = self.get_font(i[0][0])
                text_tmp.extend(([i[0][0][1], paragraph_type]))
                self.lines.append(text_tmp)
                text_tmp = ['', 0]
                end_flg = '\n'

                document += i[1] + end_flg
            # print(document)
        return document
    
    def filltexts(self):
        # for key in img_dict.keys():
            # print(key, img_dict[key])
        #  对应框选值填入
        img_dict = self.img_dict
        line_infos = self.line_infos

        line_lack = []
        img_pos = np.zeros((self.image.shape[0], self.image.shape[1], 3), dtype=np.uint8)
        img_pos += 255
        for line_info in line_infos:
            # print(line_info)
            pos = line_info[0][0]
            center = (0.5*(pos[0]+pos[2]), 0.5*(pos[1]+pos[5]))
            cv2.rectangle(img_pos, (pos[0], pos[1]), (pos[6], pos[7]), (255, 0, 0), 1)
            lack_flg = True
            for key, tp in img_dict.items():
                for t in tp:
                    if t[0][0] < center[0] < t[0][0]+t[0][2] and t[0][1] < center[1] < t[0][1]+t[0][3]:
                        t[1].append(line_info)
                        lack_flg = False
                        break
            if lack_flg:
                line_lack.append(line_info)

        self.line_lack = line_lack

        print(img_dict['table'])
        tables = img_dict['table']  # 格式示例 [[[195, 1343, 1448, 946], []]]
        self_tables = []
        for table in tables:
            table_img = self.image[table[0][1]-20:table[0][1]+table[0][3]+20, table[0][0]-20:table[0][0]+table[0][2]+20]
            try:
                # Image.fromarray(table_img).show()
                retval, buffer = cv2.imencode('.jpg', table_img)
                pic_str = base64.b64encode(buffer)
                pic_str = pic_str.decode()
                r = requests.post(TABLE_URL, data={"img": pic_str})
                img_byte = base64.b64decode(r.content.decode("utf-8"))
                img_np_arr = np.fromstring(img_byte, np.uint8)
                image = cv2.imdecode(img_np_arr, cv2.IMREAD_COLOR)
                image = cv2.resize(image, (table_img.shape[1], table_img.shape[0]))
                # Image.fromarray(image).show()
                Image.fromarray(image).save('layout_unet/table.jpg')
                table_infos = extract_table(image, table_img)
                if table_infos != 'not table' and table_infos:
                    info = table_infos[0]
                    print(11111)
                    self.lines.append([[generate_table(info, table_img), 'table', 10, 1, 1], 10, table[0][1], 'table'])
                    # print(a)
                    # print(111, self.tables)
            except Exception as ex:
                print(ex)
                pass
        print(self_tables)
        # self.tables = self_tables
        # self.lines.append(se)
            # print(table_infos)
            
        # print(111, line_lack)
        return img_dict
    
    def predict(self):
        result_temp = self.filltexts()
        line_lack = self.line_lack
        imgh, imgw = self.image.shape[:-1]
        # 将没有对应unet框的区域进行填入
        for line in line_lack:
            if line[0][0][1] < 0.2*imgh:
                x, y, w, h = line[0][0][0], line[0][0][1], line[0][0][2]-line[0][0][0], line[0][0][5]-line[0][0][1]
                result_temp['header'].append(([x, y, w, h], [line]))
            elif line[0][0][5] > 0.8*imgh:
                x, y, w, h = line[0][0][0], line[0][0][1], line[0][0][2]-line[0][0][0], line[0][0][5]-line[0][0][1]
                result_temp['footer'].append(([x, y, w, h], [line]))
            else:
                x, y, w, h = line[0][0][0], line[0][0][1], line[0][0][2]-line[0][0][0], line[0][0][5]-line[0][0][1]
                result_temp['text'].append(([x, y, w, h], [line]))
                print(11111111111)
        self.result_temp = result_temp


        # 还原页眉页脚， 字体大小
        img_dict = self.result_temp
        for header in img_dict['header']:
            res = self.layout(header[1], header[0][2], 'header')
            # print('header', header)

        for footer in img_dict['footer']:
            res = self.layout(footer[1], footer[0][2], 'footer')
            # print('footer', res)

        # 还原段落， 段落字体大小
        for texts in img_dict['text']:
            # for table in img_dict['table']:
            tables = [table[0] for table in img_dict['table']]
            if checkIsInTable(tables, texts):
                print('is in table')
                continue
            # print(222, texts)
            # # 将框选区域移动到最左端
            # line_x_list = [texts[0][0]]
            # for text in texts[1]:
            #     for line_bbox in text[0]:
            #         # print(2,line_bbox)
            #         for i in range(len(line_bbox)):
            #             if i in [0, 2, 4, 6]:
            #                 line_x_list.append(line_bbox[i])
            
            # min_x = min(line_x_list)
            # for index in range(len(texts[1])):
            #     line = texts[1][index]
            #     temp = [0, 2, 4, 6]
            #     texts[1][index] = [[[line[0][0][x]-min_x if x in temp else line[0][0][x] for x in range(len(line[0][0]))],
            #     [line[0][0][x]-min_x if x in temp else line[0][0][x] for x in range(len(line[0][0]))]], line[1]]
                # print(texts[1][index])
            # print(texts[1], texts[0][2])
            res = self.layout(texts[1], texts[0][2], 'text')
            # print(res)
        return res

    def save2docx(self):
        self.predict()
        res = self.lines
        res = sorted(res, key=lambda x: x[2])
        # print(res)
        doc = self.document
        
        sec = doc.sections[-1]
        header = sec.header
        header.is_linked_to_previous = False
        footer = sec.footer
        footer.is_linked_to_previous = False
        paragraph_type_dict = {'left': WD_PARAGRAPH_ALIGNMENT.LEFT, 'center': WD_PARAGRAPH_ALIGNMENT.CENTER, 'right': WD_PARAGRAPH_ALIGNMENT.RIGHT}
        # print(1111, res)
        print(paragraph_type_dict)
        
        # print(self.tables)
        # for i in self.tables:
        #     doc = restore_table(doc, i, Image.fromarray(self.image))
        for i in res:
            print(1111, i)
            if 'table' in i[-1]:
                doc = restore_table(doc, i[0], Image.fromarray(self.image))
            header.paragraphs[-1].style.font.size = Pt(9)
            if 'header' in i[-1]:
                header.paragraphs[-1].text += i[0]
                if '_' in i[-1]:
                    header.paragraphs[-1].alignment = paragraph_type_dict[i[-1].split('_')[1]]
                continue
            if'footer' in i[-1]:
                footer.paragraphs[-1].text += i[0]
                if '_' in i[-1]:
                    footer.paragraphs[-1].alignment = paragraph_type_dict[i[-1].split('_')[1]]
                continue
            if 'table' in i[-1]:
                # # TODO
                continue
            if 'text' in i[-1]:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.first_line_indent = Inches(0.3)
                paragraph.style.font.size = Pt(i[1])
                runner = paragraph.add_run(i[0])
                runner.font.size = Pt(i[1])
            if '_' in i[-1]:
                paragraph.paragraph_format.alignment = paragraph_type_dict[i[-1].split('_')[1]]
        return doc


if __name__ == "__main__":
    img = Image.open(r'origin.jpg')
    # img.thumbnail((1500, 1500), Image.ANTIALIAS)
    img = np.array(img)

    with open('pos.pkl', 'rb') as f:
        line_infos = pickle.load(f)
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # type_pos = unetpic2typepos(img)
    # type_pos = filltexts(type_pos, line_infos)
    unetparse = unetParse(img, line_infos, document)
    res = unetparse.save2docx()
    res.save('test.docx')