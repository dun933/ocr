import math
import os
import re

from font import get_detail

from PIL import Image
import numpy as np
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from add_picture_float import add_picture_float
from restore_table import restore_table


def get_blank1(y, h, dt):
    if h < 35:
        h = 40
    table_ = [i[1][3] for i in dt if i[0] == 'table']
    for j in dt:
        if j[0] != 'table':
            if y < j[0][0][5]:
                blanks = int((j[0][0][5]-y)//h-1)
                if blanks > 15:
                    blanks = int(blanks/3)
                elif blanks > 10:
                    blanks = int(blanks/2)
                return blanks
            if table_:
                for t in table_:
                    if t[1] < y < t[1]+t[3]:
                        return 0
        else:
            if y < j[1][3][1]:
                blanks = int((j[1][3][1] - y) // h - 1)
                if blanks > 10:
                    blanks = int(blanks/2)
                return blanks
    return 0


def row_space(y_, font_size, dt):
    if font_size == 9:
        font_ = 20
    elif font_size == 10.5:
        font_ = 35
    elif font_size == 12:
        font_ = 50
    else:
        font_ = 65
    # font_ /= 1
    dt = [i for i in dt if i[0] != 'table']
    dt = sorted(dt, key=lambda x: x[0][0][5])
    for index, i in enumerate(dt):
        if y_ == i[0][0][5]:
            if index == 0:
                return 0
            else:
                space = dt[index][0][0][1] - dt[index-1][0][0][5]
                try:
                    if font_ / space < 2:
                        return 1
                except:
                    pass
    return 0


def get_font1(dt, x_y_list):
    if x_y_list:
        target_x_y = sorted(x_y_list, key=lambda x: x[1])
        x0, y0 = dt[0][0][0], dt[0][0][1]
        y1 = dt[0][0][5]
        closed_y = list(i for i in target_x_y if y0 < (i[1]+i[3])/2 < y1)

        if closed_y:
            h = closed_y[0][3] - closed_y[0][1]
        else:
            h = dt[0][0][5] - dt[0][0][1]
    else:
        h = dt[0][0][5] - dt[0][0][1]

    if h < 30:
        return 9, h
    elif 30 <= h <= 50:
        return 10.5, h
    elif 50 < h <= 65:
        return 12, h
    elif 65 < h < 80:
        return 14, h
    return 16, h


def add_position_func_old(img, data, l, mode):
    w, h = img.size
    [x_y_list, img_list] = get_detail(img)  # 获取传统方法文本行、图片位置

    data_info = []
    if mode == 0:
        for index, i in enumerate(data):
            if i[0] == 'table':
                data_info.append([i, 'table', 10.5, 1, 0])  # get_blank1(i[1][0][1][3]
            else:
                y_ = i[0][0][5]
                font_size, h = get_font1(i, x_y_list)
                blank_size = get_blank1(y_, h, data)
                row_size = row_space(y_, font_size, data)

                # 根据y高度比判断上一段结束
                if index > 1:
                    if data[index - 1][0] != 'table' and data[index - 2][0] != 'table':
                        y_1 = i[0][0][1] - data[index - 1][0][0][5]
                        y_2 = data[index - 1][0][0][1] - data[index - 2][0][0][5]
                        try:
                            if y_1 / y_2 > 1.5 and data_info[-1][1] == 'general':
                                print(data_info[-1][1])
                                # print('11111111111111', i)
                                data_info[-1][1] = 'sub'
                        except:
                            pass

                if i[0][0][0] > 0.5 * w:  # 最左x大于一半
                    position = 'right'
                elif i[0][-1][-3] < 0.80 * w and i[0][0][0] > 0.20 * w and 0.9 * w < i[0][0][0] + i[0][-1][-
                    3] < 1.1 * w:  # 居中判断
                    position = 'center'
                    # print(position, i)
                elif i[0][0][0] > 0.35 * w and i[0][-1][-3] > 0.7 * w:  # 居右判断
                    position = 'right'
                elif i[0][-1][-3] < 0.8 * w:
                    position = 'sub'
                else:
                    position = 'general'
                data_info.append([i, position, font_size, blank_size, row_size])
    else:
        for index, i in enumerate(data):
            if i[0] == 'table':
                data_info.append([i, 'table', 10.5, 1, 0])  # get_blank1(i[1][0][1][3]
            else:
                y_ = i[0][0][5]
                h_ = i[0][0][5] - i[0][0][1]
                font_size, h = get_font1(i, x_y_list)
                blank_size = get_blank1(y_, h, data)
                row_size = row_space(y_, font_size, data)

                # 根据y高度比判断上一段结束
                if index > 1:
                    if data[index - 1][0] != 'table' and data[index - 2][0] != 'table':
                        y_1 = i[0][0][1] - data[index - 1][0][0][5]
                        y_2 = data[index - 1][0][0][1] - data[index - 2][0][0][5]
                        try:
                            if y_1 / y_2 > 1.5 and data_info[-1][1] == 'general':
                                data_info[-1][1] = 'sub'
                        except:
                            pass

                if i[0][0][0] > 0.5*w:  # 最左x大于一半
                    position = 'right'
                elif i[0][-1][-3] < 0.80*w and i[0][0][0] > 0.20*w and 0.9*w < i[0][0][0] + i[0][-1][-3] < 1.1*w:  # 居中判断
                    position = 'center'
                elif l[2] - i[0][-1][-3] > 1.75*h_ and i[0][0][0] - l[0] > 1.75*h_ and abs(i[0][0][0]-l[0]-l[2] + i[0][-1][-3]) < 0.05*w:
                    position = 'center'
                elif i[0][0][0] > 0.35 * w and i[0][-1][-3] > 0.7 * w:  # 居右判断
                    position = 'right'
                elif i[0][-1][-3] < 0.8*w:
                    position = 'sub'
                elif i[0][0][0] - l[0] > 1.75*h_ and index != 0:
                    if data_info[-1][1] == 'general':
                        data_info[-1][1] = 'sub'
                    position = 'general'
                else:
                    position = 'general'
                data_info.append([i, position, font_size, blank_size, row_size])

    # 判断图片对象位置、大小
    if img_list:
        for im in img_list:
            pos = 'center'
            size = 0.1
            flg = False
            for index, dt in enumerate(data_info[::-1]):
                if dt[1] == 'table':
                    dt_y = dt[0][1][3][1]
                elif dt[1] == 'picture':
                    dt_y = dt[0][0][1]
                else:
                    dt_y = dt[0][0][0][1]
                if dt_y < im[1]:
                    if im[0] > 0.5 * w:
                        pos = 'right'
                    elif im[2] < 0.5 * w:
                        pos = 'left'
                    else:
                        pos = 'center'

                    size = 0.7 * w / (im[2] - im[0])
                    if size > 3:
                        size = 3
                    data_info.insert(len(data_info) - index, [[im, ''], 'picture', pos, size, 0])
                    flg = True
                    break
            if not flg:
                data_info.append([[im, ''], 'picture', pos, size, 0])
    return data_info  # , img_list


def add_position_func(img, data, l, mode):
    print('mode:', mode)
    w, h = img.size
    [x_y_list, img_list] = get_detail(img)  # 获取传统方法文本行、图片位置

    # print(l)
    if mode == 1:
        if abs(l[0]+l[2]) > l[5]+l[4]:
            w = l[5]+l[4]
        else:
            w = l[0]+l[2]
    # import cv2
    # imm = np.array(img)
    # imm = cv2.line(imm, (int((l[2] - l[0])), 0), (int((l[2] - l[0]) / 0.8), h), (255, 0, 0), 1)
    # cv2.rectangle(imm, (l[0], l[1]), (l[2], l[3]), (0, 255, 0), 1)
    # print('1111111111111111111111')
    # Image.fromarray(imm).show()

    # 判断图片对象位置、大小
    img_pos_list = []
    if img_list:
        for im in img_list:
            if im[0] > 0.5 * w:
                pos = 'right'
            elif im[2] < 0.5 * w:
                pos = 'left'
            else:
                pos = 'center'
            size = 0.7 * w / (im[2] - im[0])
            if size > 3:
                size = 3
            img_pos_list.append([[im, ''], 'picture', pos, size, 0])

    if mode == 3:
        return [[], img_pos_list]

    data_info = []

    h_list = [x[0][0][5] - x[0][0][1] for x in data if x[0] != 'table']
    # max_x = sorted(data, key=lambda x: x[0][0][2] - x[0][0][0] if x[0] != 'table' else x[1][3][2])
    data_length = len(data)

    for index, i in enumerate(data):
        if i[0] == 'table':
            data_info.append([i, 'table', 10.5, 1, 0])  # get_blank1(i[1][0][1][3]
        else:
            y_ = i[0][0][5]
            h_ = i[0][0][5] - i[0][0][1]
            font_size, h = get_font1(i, x_y_list)
            blank_size = get_blank1(y_, h, data)
            row_size = row_space(y_, font_size, data)

            # 根据y高度比判断上一段结束
            if index > 1:
                if data[index - 1][0] != 'table' and data[index - 2][0] != 'table':
                    y_1 = i[0][0][1] - data[index - 1][0][0][5]
                    y_2 = data[index - 1][0][0][1] - data[index - 2][0][0][5]
                    y = i[0][0][5] - i[0][0][1]
                    try:
                        if y_1 / y_2 > 1.6 and data_info[-1][1] == 'general' and y < y_1:
                            print('高度比判断段落结束', data_info[-1], y_1/y_2)
                            data_info[-1][1] = 'sub'
                    except:
                        pass

            if mode == 0:  # 正常提取情况
                if i[0][0][0] > 0.5 * w:  # 最左x大于一半
                    position = 'right'
                elif i[0][0][0] > 0.4*w and i[0][0][2] < 0.6*w:
                    position = 'center'
                elif i[0][-1][-3] < 0.8 * w and i[0][0][0] > 0.2 * w and 0.9 * w < i[0][0][0] + i[0][-1][-3] < 1.1 * w:  # 居中判断
                    position = 'center'
                elif h_ > np.mean(h_list)+5 and 0.9*w < (i[0][0][2] + i[0][0][0]) < 1.1*w and index < 0.5*data_length:
                    position = 'center'
                elif i[0][0][0] > 0.35 * w and i[0][-1][-3] > 0.7 * w:  # 居右判断
                    position = 'right'
                elif i[0][-1][-3] < 0.8 * w:
                    position = 'sub'
                else:
                    position = 'general'

                if index != 0:
                    if position == 'center':
                        if re.search('。', i[1]) or len(i[1]) > 20:
                            position = 'general'

                    if data_info[-1][1] == 'general' and position != 'sub':
                        if abs(data[index - 1][0][0][0] - i[0][0][0]) < (i[0][0][5] - i[0][0][1]):
                            position = 'general'
                data_info.append([i, position, font_size, blank_size, row_size])

            else:  # w文本只占图像小部分的情况
                # w = l[0] + l[2]
                h_ = i[0][0][5] - i[0][0][1]
                # print('H_:', h_, 'h_list:', h_list)
                position = 'general'
                if i[0][0][0] > 0.5 * w:  # 最左x大于一半
                    position = 'right'
                # elif i[0][-1][-3] < 0.70 * w and i[0][0][0] > 0.30 * w and 0.9 * w < i[0][0][0] + i[0][-1][
                #     -3] < 1.1 * w:  # 居中判断
                #     position = 'center'
                elif i[0][0][-3] < 0.70 * w and i[0][0][0] > 0.30 * w and 0.9 * w < i[0][0][0] + i[0][0][
                    -3] < 1.1 * w:  # 居中判断
                    position = 'center'
                # elif l[2] - i[0][-1][-3] > 1.75 * h_ and i[0][0][0] - l[0] > 1.75 * h_ and abs(
                #         i[0][0][0] - l[0] - l[2] + i[0][-1][-3]) < 0.05 * w:
                #     position = 'center'
                elif l[2] - i[0][0][-3] > 1.75 * h_ and i[0][0][0] - l[0] > 1.75 * h_ and abs(
                        i[0][0][0] - l[0] - l[2] + i[0][0][-3]) < 0.05 * w:
                    position = 'center'
                elif h_ > np.mean(h_list) and 0.9*w < i[0][0][2] + i[0][0][0] < 1.1*w and index < 0.5*data_length:
                    position = 'center'
                elif i[0][0][0] > 0.35 * w and i[0][-1][-3] > 0.7 * w:  # 居右判断
                    position = 'right'
                elif i[0][-1][-3] < 0.8 * w:
                    position = 'sub'
                # elif i[0][0][0] - l[0] > 1.75 * h_ and index != 0:
                #     if data_info[-1][1] == 'general':
                #         data_info[-1][1] = 'sub'
                #     position = 'general'
                else:
                    position = 'general'
                # print(1111111111111111, i)
                if index != 0:
                    if position == 'center':
                        if re.search('。', i[1]) or len(i[1]) > 20:
                            position = 'general'

                    if data_info[-1][1] == 'general' and position != 'sub':
                        if abs(data[index - 1][0][0][0] - i[0][0][0]) < (i[0][0][5] - i[0][0][1]):
                            position = 'general'
                data_info.append([i, position, font_size, blank_size, row_size])
    # for i in data_info:
    #     print(i)
    return data_info, img_pos_list


def sort_paragraph(img, data):
    if not data:
        data_info, img_pos_list = add_position_func(img, data, 0, 3)

        new_img_pos_list = []
        if img_pos_list:
            for i in img_pos_list:
                new_img_pos_list.append([i[0][0], i[1], i[2], i[3], i[4]])
        return [[], new_img_pos_list]

    w, h = img.size
    # print('img_size:', img.size)
    # 最大宽度框
    max_x = sorted(data, key=lambda x: x[0][0][2] - x[0][0][0] if x[0] != 'table' else x[1][3][2])  # 添加具体内容table

    x_list = [[x[0][0][0], x[0][0][2]] if x[0] != 'table' else [x[1][3][0], x[1][3][0]+x[1][3][2]] for x in data]
    x_list = sum(x_list, [])
    x_list = [min(x_list), max(x_list)]

    if max_x[-1][0] == 'table':
        x_length = max_x[-1][1][3]  # max_x[-1][1][0][1]
        x_length = [x_length[0], x_length[1], x_length[0]+x_length[2], x_length[1]+x_length[3]]
    else:
        x_length = max_x[-1][0][0]
        x_length = [x_length[0], x_length[1], x_length[6], x_length[7]]

    x_length.extend(x_list)
    mode = 0
    if x_length[2]-x_length[0] < 0.7*w:  # 图片只占中间部分的情况
        mode = 1
    data_info, img_pos_list = add_position_func(img, data, x_length, mode)

    new_img_pos_list = []
    if img_pos_list:
        for i in img_pos_list:
            new_img_pos_list.append([i[0][0], i[1], i[2], i[3], i[4]])

    result = []
    for index, i in enumerate(data_info):
        if index == 0:
            if i[1] == 'center':  # 判断页眉
                if i[0][0][0][1] < 0.05 * h:
                    result.append([i[0][1], 'header_center', i[2], i[3], i[4]])  # header
                    continue
                result.append([i[0][1], 'center', i[2], i[3], i[4]])
                continue
            if i[1] == 'right':  # 判断页眉
                if i[0][0][0][1] < 0.05 * h:
                    result.append([i[0][1], 'header_right', i[2], i[3], i[4]])  # header
                    continue
                result.append([i[0][1], 'right', i[2], i[3], i[4]])
                continue
            if i[1] == 'table':
                result.append([i[0][1], i[1], i[2], i[3], i[4]])
                continue
            elif 0.5*w < i[0][0][-1][0]:  # i[0][0][0][3] <
                separate = i[0][1].split(' ')
                if len(separate) == 2:
                    result.append([[separate[0], separate[-1]], 'header_separate', i[2], i[3], i[4]])
                    continue
            result.append([i[0][1], i[1], i[2], i[3], i[4]])
        elif index == len(data_info)-1:
            if i[1] in ['center']:  # , 'right'  # 判断页脚
                # print(i)
                if re.compile('([0-9一二三四五六七八九十页IV])').search(i[0][1]) and i[0][0][-1][2]-i[0][0][0][0] < 0.1*w:
                    result[-1][3] = 0  # 设置blank为0
                    result.append([i[0][1], f'footer_{i[1]}', i[2], i[3], i[4]])  # 设置页脚
                    continue
            if result[-1][1] == 'general' and i[1] in ['sub', 'general']:
                result[-1][3] = 0  # 设置blank为0
                result[-1][1] = 'sub'
                result[-1][0] += i[0][1]
                continue
            else:
                if i[1] == 'general':
                    i[1] == 'sub'
                result.append([i[0][1], i[1], i[2], i[3], i[4]])
        else:
            i_content = i[0][1]
            new_i = [i[0][1], i[1], i[2], i[3], i[4]]
            if i[1] not in ['table', 'picture'] and data_info[index-1][1] not in ['table', 'picture']:
                if re.compile('(^([0-9]{1,2}|[一二三四五六七八九十]{1,3})[、.])|(^[A-Z][）)])').search(i_content) or (
                        re.search(r'(^第.{1,3}条)|(^[(（][0-9一二三四五六七八九十]{1,3}[）)])', i_content)):  # and i[0][0][0][0] > data_info[index-1][0][0][0][0])
                    if data_info[index-1][1] == 'general':
                        result[-1][1] = 'sub'
                    result.append(new_i)
                    continue
            if i[1] in ['table', 'center', 'right']:
                if result[-1][1] == 'general':  # if data_info[index-1][1] == 'general':
                    result[-1][1] = 'sub'
                result.append(new_i)
            elif result[-1][1] in ['table', 'center', 'right', 'sub', 'picture', 'header_right', 'header_center',
                                   'header_separate']:
                result.append(new_i)
            else:
                result[-1][0] += i[0][1]
                result[-1][1] = i[1]
                result[-1][2] = min([result[-1][2], i[2]])
                result[-1][3] = i[3]

    # 居中或居右有空格则改为下一行
    for index, i in enumerate(result):
        if i[1] in ['center', 'right']:
            if ' ' in i[0]:
                text = i[0]
                result[index][0] = text.split(' ')[0]
                result[index][1] = 'center'
                insert_i = i.copy()
                insert_i[0], insert_i[1] = text.split(' ')[1], 'right'
                result.insert(index+1, insert_i)
    return [result, new_img_pos_list]


def save2docx(doc, result, img, jpg_page):

    img_pos_list = result[1]
    result = result[0]

    space_after_code = 5
    line_spacing_rule_code = 1
    if len([pg for pg in result if pg[1] in ['sub', 'right', 'picture', 'center']]) >= 15:
        line_spacing_rule_code, space_after_code = 0, 2

    row_word_dict = {9: 48, 10.5: 41, 12: 36, 14: 30, 16: 27}
    # 计算文本段落及表格行数
    rows = sum([math.ceil(1.2 * i[0][-3]) if i[1] == 'table' else math.ceil((len(i[0]) / row_word_dict[i[2]])) for i
                in result if i[1] in ['center', 'general', 'right', 'sub', 'table']])
    # 计算空行数
    spac_afters_list = sum([i[3] if i[3] > 0 else 0 for i in result])
    # print('space_after', spac_afters_list)
    # print(rows)
    lines = rows + spac_afters_list
    if 30 < lines < 35:
        for i in range(len(result)):
            result[i][2] = result[i][2] - 1
    elif 35 <= lines < 40:
        for i in range(len(result)):
            result[i][2] = result[i][2] - 2
    elif lines >= 40:
        for i in range(len(result)):
            result[i][2] = result[i][2] - 3

    sec = doc.sections[-1]
    header = sec.header
    header.is_linked_to_previous = False
    footer = sec.footer
    footer.is_linked_to_previous = False
    for index, i in enumerate(result):
        if i[1].startswith('header'):
            header.paragraphs[-1].style.font.size = Pt(9)
            if i[1].endswith('separate'):
                header.paragraphs[-1].text = i[0][0] + '\t' + i[0][1]
                header.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                continue
            if i[1].endswith('center'):
                header.paragraphs[-1].text = i[0]
                header.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                continue
            header.paragraphs[-1].text = i[0]
            header.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif i[1].startswith('footer'):
            footer.paragraphs[-1].style.font.size = Pt(8)
            footer.paragraphs[-1].text = i[0]
            footer.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif i[1] == 'table':
            doc = restore_table(doc, i, img)
            if index == len(result)-1 or index == len(result)-2:
                doc.add_paragraph()
        else:
            if i[1] == 'center':
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.style.font.size = Pt(i[2])  #
                runner = paragraph.add_run(i[0])
                runner.font.size = Pt(i[2])
                runner.bold = True
            elif i[1] == 'right':
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                paragraph.style.font.size = Pt(i[2])
                runner = paragraph.add_run(i[0])
                runner.font.size = Pt(i[2])
            else:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.first_line_indent = Inches(0.3)
                paragraph.style.font.size = Pt(i[2])
                runner = paragraph.add_run(i[0])
                runner.font.size = Pt(i[2])
            if i[3] > 0:
                for j in range(i[3]):
                    blank_paragraph = doc.add_paragraph('')
                    blank_paragraph.style.font.size = Pt(7)
                    blank_paragraph.paragraph_format.space_after = Pt(0)
                    blank_paragraph.paragraph_format.line_spacing_rule = Pt(0)
            if i[4] == 0:
                paragraph.paragraph_format.line_spacing_rule = 0
            if i[3] > 0:
                space_after_code = 0
            paragraph.paragraph_format.space_after = Pt(space_after_code)
            paragraph.paragraph_format.line_spacing_rule = line_spacing_rule_code  # 1

    # 写入图片
    if img_pos_list:
        data = [i for i in result if i[1] in ['general', 'sub', 'table', 'center']]
        if not data:
            print('result为空')
            doc.add_paragraph()
        for i in img_pos_list:
            try:
                Image.fromarray(np.array(img)[i[0][1]:i[0][3], i[0][0]:i[0][2]]).save('img_s.jpg')
                if i[0][0] > img.width / 2:
                    pos_w = 7534275 * (i[0][0] / img.width)
                else:
                    pos_w = 7534275 * (i[0][0] / img.width)
                if i[0][1] < img.width / 2:
                    pos_y = 10592435 * (i[0][1] / img.height)
                else:
                    pos_y = 10592435 * (2 * i[0][1] - i[0][3]) / img.height
                pos = f'{i[0][0]}_{i[0][1]}_{i[0][2]}_{i[0][3]}_{jpg_page}_{img.width}_{img.height}'
                add_picture_float(doc, 'img_s.jpg', (pos_w, pos_y), pos, scales=5)
                os.remove('img_s.jpg')
            except:
                pass

    # doc.add_section()
    return doc


if __name__ == '__main__':
    import pickle
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    base_path = r'C:\Users\Admin\Desktop\新建文件夹'
    n = 0
    for i in os.listdir(base_path):
        if n > 10:
            break
        n+=1
        if i.endswith('pkl'):
            name = f'{base_path}/{i}'
            img = Image.open(name.replace('.pkl', '.jpg'))
            img.thumbnail((2500, 2500), Image.ANTIALIAS)
            data = pickle.load(open(name, 'rb'))
            result = sort_paragraph(img, data)
            print(result)
            doc = save2docx(doc, result, img, 0)
    doc.save("aaaa.docx")
