import pickle

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt, Inches


# 设置表格行高
def set_row_height(table, begin, end, ratio):
    # 表格元素、起始行、终止行、行占原图比例
    try:
        for row in range(begin, end):
            row = table.rows[row]
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(int(ratio*13000)))  # 13000为在纸张21.5*27.9cm, 上下距离25.4mm时页面设置最大值

            # trHeight.set(qn('w:val'), str(30))  # 强制最小 根据cell内容自适应
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)
    except Exception as ex:
        print('set_row_height', ex)


# 设置表格列宽
def set_column_width(table, begin, end, width):
    try:
        for col in range(begin, end):
            # WPS设置列宽
            table.columns[col].width = width
            # office2016设置列宽
            rows = table.rows
            for row in range(len(rows)):
                table.cell(row, col).width = width
    except Exception as ex:
        print('set_column_width', ex)


def restore_table(doc, i, img):
    # import pickle
    # pickle.dump(i, open('1.pkl', 'wb'))
    i = ['table', i[0]]
    table_row, table_col = i[1][1], i[1][2]
    table = doc.add_table(i[1][1], i[1][2], style='Table Grid')

    # 表格居中
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for j in i[1][0]:
        # print(j)
        try:
            d = j[0]
            row_begin, col_begin = d['row_begin'], d['col_begin']
            row_end, col_end = d['row_end'], d['col_end']

            # 排除错误情况
            if row_end == 0 or col_end == 0 or row_begin == table_row or col_begin == table_col:
                continue
            if row_end == row_begin:
                row_end += 1
            if col_end == col_begin:
                col_end += 1

            # 设置行高
            if col_begin == 0:
                rs = row_end - row_begin
                row_size = j[1][2]/(img.height*rs)/0.8
                # print('row', row_size)
                set_row_height(table, row_begin, row_end, row_size)

            # 设置列宽
            if row_begin == 0:
                cs = col_end - col_begin
                col_size = 10 * (j[1][1] / cs) / img.width
                # print('col', col_size)
                set_column_width(table, col_begin, col_end, Inches(col_size))

            # 向单元格中添加值
            try:
                cell = table.cell(row_begin, col_begin)
                run = cell.paragraphs[0].add_run(j[1][0])
                cell.merge(table.cell(row_end-1, col_end-1))

                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                run.font.size = Pt(7)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                cell.paragraphs[0].paragraph_format.line_spacing_rule = 0
                cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            except Exception as ex:
                print('cell_error', ex, j)
        except Exception as ex:
            print('table_error', ex, j)
    return doc


if __name__ == '__main__':
    img = Image.open(r'C:\Users\Admin\Desktop\texts.jpg')
    img.thumbnail((2500, 2500))
    doc = Document()
    data = pickle.load(open(r'F:\paragraph_restore\pdf\p1\76\76_103.pkl', 'rb'))
    data = [i for i in data if i[0] == 'table']
    for i in data:
        doc = restore_table(doc, i, img)
    doc.save('1.docx')
    pass
