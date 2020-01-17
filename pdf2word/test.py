import os

# from extract_table_1 import extract_table
from pdf2word.pdf_table import generate_table, extract_table

os.environ['CUDA_VISIBLE_DEVICES'] = '-1'
import cv2
import fitz
import pickle

import numpy as np
from io import open
from PIL import Image
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
# from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter, PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBoxHorizontal, LTImage, LTFigure, LTRect, LTLine
from pdfminer.pdfparser import PDFDocument, PDFParser

# from pdf_word.pdf_table import extract_table, generate_table
from pdf2word.layout_test_pdf import sort_paragraph, save2docx


# def pdf_pages(file_path):
#     with open(file_path, 'rb') as file:
#         praser = PDFParser(file)
#         # 创建一个PDF文档
#         doc = PDFDocument()
#         # 连接分析器 与文档对象
#         praser.set_document(doc)
#         doc.set_parser(praser)
#         # 提供初始化密码
#         # 如果没有密码 就创建一个空的字符串
#         doc.initialize()
#         # print('is_extractable', doc.is_extractable)
#         # 创建PDf 资源管理器 来管理共享资源
#         rsrcmgr = PDFResourceManager()
#         # 创建一个PDF设备对象
#         laparams = LAParams()
#         device = PDFPageAggregator(rsrcmgr, laparams=laparams)
#         # 创建一个PDF解释器对象
#         interpreter = PDFPageInterpreter(rsrcmgr, device)
#
#         page = list(doc.get_pages())
#         return len(page)


def pdf_is_text(file_path):
    # TODO
    return False
    try:
        with open(file_path, 'rb') as file:
            praser = PDFParser(file)
            # 创建一个PDF文档
            doc = PDFDocument()
            # 连接分析器 与文档对象
            praser.set_document(doc)
            doc.set_parser(praser)
            # 提供初始化密码
            # 如果没有密码 就创建一个空的字符串
            doc.initialize()

            # 判断是否加密
            if doc.encryption:
                # pdf = fitz.Document(file_path)
                # pdf.save('fitz_decrypt.pdf')
                with open('fitz_decrypt.pdf', 'rb') as f:
                    praser = PDFParser(f)
                    doc = PDFDocument()
                    praser.set_document(doc)
                    doc.set_parser(praser)
                    doc.initialize()

            # print('is_extractable', doc.is_extractable)
            # 创建PDf 资源管理器 来管理共享资源
            rsrcmgr = PDFResourceManager()
            # 创建一个PDF设备对象
            laparams = LAParams()
            device = PDFPageAggregator(rsrcmgr, laparams=laparams)
            # 创建一个PDF解释器对象
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            # 循环遍历列表，每次处理一个page的内容
            first_three = [0, 0, 0]
            for index, page in enumerate(doc.get_pages()):  # doc.get_pages() 获取page列表
                if index < 3:
                    interpreter.process_page(page)
                    # 接受该页面的LTPage对象
                    layout = device.get_result()
                    for i in layout:
                        if isinstance(i, LTTextBoxHorizontal):
                            print(i.get_text())
                            first_three[index] += len(i.get_text())
                else:
                    break
            # print(first_three)
            # 如果前三页字数相同且少于50则判断为图片类PDF
            if max(first_three) < 50:  # == min(first_three) and first_three[0]
                return False
            else:
                return True
    except Exception as ex:
        return False


def read_from_pdf(file_path, page):
    with open(file_path, 'rb') as file:
        praser = PDFParser(file)
        # 创建一个PDF文档
        doc = PDFDocument()
        # 连接分析器 与文档对象
        praser.set_document(doc)
        doc.set_parser(praser)
        # 提供初始化密码
        # 如果没有密码 就创建一个空的字符串
        doc.initialize()

        # 判断是否加密
        if doc.encryption:
            pdf = fitz.Document(file_path)
            pdf.save('fitz_decrypt.pdf')
            with open('fitz_decrypt.pdf', 'rb') as f:
                praser = PDFParser(f)
                doc = PDFDocument()
                praser.set_document(doc)
                doc.set_parser(praser)
                doc.initialize()

        # print('is_extractable', doc.is_extractable)
        # 创建PDf 资源管理器 来管理共享资源
        rsrcmgr = PDFResourceManager()
        # 创建一个PDF设备对象
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        page = list(doc.get_pages())[page]

        interpreter.process_page(page)
        # 接受该页面的LTPage对象
        layout = device.get_result()
        result_data = []
        # data_all = [[], []]
        im_white = np.zeros((842, 595, 3), np.uint8)
        im_white = 255 - im_white
        img_data = []
        for x in layout:
            pos = [int(x.x0), int(842.25 - x.y0), int(x.x1), int(842.25 - x.y0), int(x.x0), int(842.25 - x.y1),
                   int(x.x1), int(842.25 - x.y1), 0]
            if isinstance(x, LTRect):
                cv2.rectangle(im_white, (int(x.x0), int(842.25 - x.y0)), (int(x.x1), int(842.25 - x.y1)), (255, 0, 0), 2)
            if isinstance(x, LTImage):
                img_data.append([pos[0], pos[1], pos[6], pos[7]])
            elif isinstance(x, LTFigure):
                img_data.append([pos[0], pos[1], pos[6], pos[7]])
            elif isinstance(x, LTTextBoxHorizontal):
                # print(x.get_text())
                for obj in x._objs:
                    for i in obj:
                        try:
                            x0, y0, x1, y1 = int(i.x0), int(842.25 - i.y0), int(i.x1), int(842.25 - i.y1)
                            result_data.append([i._text, [x0, y0, x1, y1]])
                        except:
                            pass
    return result_data, img_data, im_white  # data_all,


def extract_pdf_text(pdf_path, page):
    data, im_data, im_white = read_from_pdf(pdf_path, page)
    data = sorted(data, key=lambda x: x[1][1])

    pdf = fitz.open(pdf_path)
    page_fitz = pdf[page]
    trans = fitz.Matrix(3, 3).preRotate(0)
    pm = page_fitz.getPixmap(matrix=trans, alpha=False)
    img = Image.frombytes("RGB", [pm.width, pm.height], pm.samples)
    if img.height > img.width:
        img.thumbnail((595.5, 842.25), Image.ANTIALIAS)
    else:
        img.thumbnail((842.25, 595.5), Image.ANTIALIAS)

    tables = extract_table(np.array(img))
    if tables == 'not table':
        tables = []

    table_data = []
    for i in data:
        if tables:
            cen_point = 0.5 * (i[1][1] + i[1][3])
            flg = False
            for table in tables:
                table_pos = table[0][-1]
                if table_pos[1] < cen_point < table_pos[1] + table_pos[3]:
                    flg = True
            if flg:
                table_data.append(i)

    text_data = [i for i in data if i not in table_data]

    new_text_data = []
    for index, i in enumerate(text_data):
        if index == 0:
            new_text_data.append([i])
        else:
            if i[1][1] - new_text_data[-1][-1][1][1] < 2:
                new_text_data[-1].append(i)
            else:
                new_text_data.append([i])

    new_text_data = [sorted(i, key=lambda x: x[1][0]) for i in new_text_data]
    new_text_data = [[''.join([j[0] for j in i]), [i[0][1][0], i[0][1][1], i[-1][1][2], i[0][1][1], i[0][1][0],
                                                   i[0][1][3], i[-1][1][2], i[0][1][3], 0]] for i in new_text_data]
    new_text_data = [[[np.array(i[1]), np.array(i[1])], i[0]] for i in new_text_data]

    for table in tables:
        table_index = 0
        for index, i in enumerate(new_text_data):
            if i[0] == 'table':
                if table[0][1][1] > i[1][3][1]:
                    table_index = index + 1
            elif table[0][1][1] > i[0][0][1]:
                table_index = index + 1
        new_text_data.insert(table_index, ['table', generate_table(table, table_data=table_data)])

    data = [new_text_data, im_data]

    result = sort_paragraph(img, data)
    return result, img


if __name__ == '__main__':
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    pdf_path = r'F:\数据\txt_pdfs\16550822.pdf'

    start_page, end_page = 0, 1
    for page in range(start_page, end_page):  # pdf_pages(pdf_path)
        print(page)
        print(pdf_is_text(pdf_path))
        result, img = extract_pdf_text(pdf_path, page)
        print(result)
        doc = save2docx(doc, result, img, page)
        if page == end_page-1:
            continue
        doc.add_section()

    doc.save('aaa2.docx')
    # for i in os.listdir(r'F:\数据\txt_pdfs'):
    #     if not pdf_is_text(r'F:\数据\txt_pdfs/' + i):
    #         print(i)

    # print(pdf_is_text(pdf_path))

