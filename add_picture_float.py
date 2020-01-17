import random

from lxml import etree
from docx import Document


def drawing_pic_xml_template(id, rId, c, position, pos, behindDoc=0):
    # print(11111111, pos)
    positionH = str(int(position[0]))
    # print('1111', positionH, type(positionH))
    # positionH = str(int(7534275*(1-positionH)))
    positionV = str(int(position[1]))
    cx = str(c[0])
    cy = str(c[1])
    id = str(id)
    rId = str(rId)
    descr = ''.join([random.choice('0123456789abcdef') for i in range(31)])
    behindDoc = str(behindDoc)
    drawing_pic_xml = '''
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
      <w:rPr>
        <w:rFonts w:hint="eastAsia"/>
        <w:lang w:val="en-US" w:eastAsia="zh-CN"/>
      </w:rPr>
      <w:drawing>
        <wp:anchor distT="0" distB="0" distL="114935" distR="114935" simplePos="0" relativeHeight="251659264" behindDoc="''' + behindDoc + '''" locked="0" layoutInCell="1" allowOverlap="1">
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="page">
            <wp:posOffset>''' + positionH + '''</wp:posOffset>
          </wp:positionH>
          <wp:positionV relativeFrom="page">
            <wp:posOffset>''' + positionV + '''</wp:posOffset>
          </wp:positionV>
          <wp:extent cx="''' + cx + '''" cy="''' + cy + '''"/>
          <wp:effectExtent l="0" t="0" r="5715" b="7620"/>
          <wp:wrapNone/>
          <wp:docPr id="''' + id + '''" name="图片 2" descr="''' + descr + '''"/>
          <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>
          </wp:cNvGraphicFramePr>
          <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
            <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
              <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:nvPicPr>
                  <pic:cNvPr id="''' + id + '''" name="图片 2" descr="''' + descr + '''"/>
                  <pic:cNvPicPr>
                    <a:picLocks noChangeAspect="1"/>
                  </pic:cNvPicPr>
                </pic:nvPicPr>
                <pic:blipFill>
                  <a:blip r:embed="''' + rId + '''" descr="''' + pos + '''"/>
                  <a:stretch>
                    <a:fillRect/>
                  </a:stretch>
                </pic:blipFill>
                <pic:spPr>
                  <a:xfrm>
                    <a:off x="0" y="0"/>
                    <a:ext cx="''' + cx + '''" cy="''' + cy + '''"/>
                  </a:xfrm>
                  <a:prstGeom prst="rect">
                    <a:avLst/>
                  </a:prstGeom>
                </pic:spPr>
              </pic:pic>
            </a:graphicData>
          </a:graphic>
        </wp:anchor>
      </w:drawing>
    </w:r>'''
    return etree.fromstring(drawing_pic_xml)


def add_picture_float(doc, img_path, position, pos, scales=10):

    new_pic_xml = doc.part.new_pic_inline(img_path, width=None, height=None)  # .image

    # print(new_pic_xml.xml)

    id = new_pic_xml.xpath('/wp:inline/wp:docPr')[0].id
    rId = new_pic_xml.xpath('/wp:inline/a:graphic/a:graphicData/pic:pic/pic:blipFill/a:blip')[0].embed
    cx = new_pic_xml.xpath('/wp:inline/wp:extent')[0].cx
    cy = new_pic_xml.xpath('/wp:inline/wp:extent')[0].cy

    ele = doc._element  # .xml

    # section =
    sdct = ele.xpath("/w:document/w:body/w:p")
    if len(sdct) == 0:
        doc.add_paragraph()
    sdct = ele.xpath("/w:document/w:body/w:p")

    c = (int(cx/scales), int(cy/scales))
    sdct[-1].insert(-1, drawing_pic_xml_template(id=id, rId=rId, c=c, position=position, pos=pos, behindDoc=1))


def main():
    path = r'C:\Users\Admin\Desktop\Desktop.docx'
    img_path = '1.jpg'
    doc = Document(path)
    position = (0, 0)
    scales = 10
    pos = [1, 2, 3, 4]
    add_picture_float(doc, img_path, position, pos,  scales)
    doc.save(r'C:\Users\Admin\Desktop\Desktop - 副本111.docx')


if __name__ == '__main__':
    main()
